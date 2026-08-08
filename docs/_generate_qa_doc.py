"""Regenerate LinkUp launch QA user test case Word documents (v2.0 sequential suites)."""
from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

ROOT = Path(__file__).resolve().parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from qa_launch_cases import (  # noqa: E402
    MOBILE_CASE_COUNT,
    MOBILE_SECTIONS,
    TOTAL_CASE_COUNT,
    WEB_CASE_COUNT,
    WEB_SECTIONS,
    QaSection,
)

DOCX_PATH = ROOT / "LinkUp-QA-User-Test-Cases.docx"
DOC_PATH = ROOT / "LinkUp-QA-User-Test-Cases.doc"
GENERATED = date.today().isoformat()
VERSION = "2.0"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_table_row(table, cells: list[str]) -> None:
    row = table.add_row().cells
    for i, value in enumerate(cells):
        row[i].text = value


def add_test_case(
    doc: Document,
    case_id: str,
    title: str,
    module: str,
    priority: str,
    preconditions: str,
    steps: list[str],
    expected: str,
) -> None:
    doc.add_heading(f"{case_id}: {title}", level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Test ID"
    hdr[1].text = "Module"
    hdr[2].text = "Priority"
    hdr[3].text = "Result"
    add_table_row(table, [case_id, module, priority, "Pass / Fail / Blocked"])
    add_para(doc, "Preconditions:", bold=True)
    add_para(doc, preconditions)
    add_para(doc, "Test Steps:", bold=True)
    for i, step in enumerate(steps, start=1):
        add_para(doc, f"{i}. {step}")
    add_para(doc, "Expected Result:", bold=True)
    add_para(doc, expected)
    doc.add_paragraph()


def render_suite(
    doc: Document,
    sections: list[QaSection],
    id_prefix: str,
    platform_label: str,
) -> int:
    """Render sequential test cases; returns number of cases written."""
    counter = 0
    for section in sections:
        add_heading(doc, section["name"], level=2)
        module = section["name"].split(". ", 1)[-1] if ". " in section["name"] else section["name"]
        for case in section["cases"]:
            counter += 1
            case_id = f"{id_prefix}-{counter:03d}"
            pre = f"Platform: {platform_label}. {case['preconditions']}"
            add_test_case(
                doc,
                case_id,
                case["title"],
                module,
                case["priority"],
                pre,
                case["steps"],
                case["expected"],
            )
    return counter


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading("LinkUp — Launch QA User Test Cases", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Quality Assurance · Sequential Manual UAT · Mobile + Web")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(
        f"Document version: {VERSION}    Generated: {GENERATED}    "
        f"Mobile: {MOBILE_CASE_COUNT} cases    Web: {WEB_CASE_COUNT} cases    Total: {TOTAL_CASE_COUNT}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. Launch Goal", 1)
    add_para(
        doc,
        "This document is the single source of truth for pre-launch user acceptance testing. "
        "The goal is a bug-free release where users can sign up, discover meetups, create and join plans, "
        "pay and receive payouts safely, confirm meetups, resolve disputes, and enjoy messaging and safety "
        "features without friction on both the LinkUp mobile app (iOS/Android) and LinkUp Web (browser).",
    )
    add_para(
        doc,
        "Execute test cases in sequential order within each part. Do not skip Critical cases. "
        "Log every failure with steps to reproduce, build/version, and screenshots before sign-off.",
    )

    add_heading(doc, "2. Scope & Platform Split", 1)
    add_para(
        doc,
        "Part A (TC-M-001 … TC-M-"
        f"{MOBILE_CASE_COUNT:03d}) covers the LinkUp mobile app only — iOS and Android builds.",
    )
    add_para(
        doc,
        "Part B (TC-W-001 … TC-W-"
        f"{WEB_CASE_COUNT:03d}) covers LinkUp Web only — desktop and mobile browsers against linkup-web.",
    )
    add_para(
        doc,
        "Part C is a cross-platform parity checklist for flows that must behave consistently on both surfaces. "
        "Mobile and web cases are listed exclusively in their respective parts; no feature area is omitted.",
    )
    add_para(
        doc,
        "Feature coverage includes: authentication, onboarding (with validation), discover and Meetr, "
        "mood plans (20 km reach), saved plans and offers, plan creation (standard, mood, group / Annexure B), "
        "negotiation and agreement, escrow and Flutterwave payments, bank transfer and refund accounts, "
        "wallet disbursement and withdrawals, meetup confirmation, exigency reports, disputes (video + chat consent), "
        "reviews and ratings, messaging, live location, push and in-app notifications, KYC, premium tiers, "
        "plan management, invitations and deep links, policy modals, support, admin tooling, and platform reliability.",
    )

    add_heading(doc, "3. Test Environment", 1)
    for line in [
        "Staging or production Supabase with Auth SMTP (signup/reset emails deliver).",
        "Mobile: physical devices for push, deep links (linkup://auth/callback), camera/KYC, and live location.",
        "Web: linkup-web deployed against the same Supabase project; Chrome + Safari; test mobile viewport (375px).",
        "Test accounts: unverified member, verified member, Silver/Gold/Platinum, group-plan host, admin (admins row).",
        "Flutterwave test keys; edge functions deployed: create-escrow-payment, flutterwave-webhook, disburse-wallet, "
        "disbursement-sweep, submit-arrival-nudge, submit-dispute-video, submit-exigency-report, verify-bank-account, etc.",
        "Migrations applied: mood_reach_and_notifications, onboarding_complete_validation, group_plan_annexure_b, "
        "meetup_confirmation_disbursement, rating_review_system.",
        "Second device/account recommended for messaging, negotiation, split escrow, live location, and dual confirmation.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "4. Priority Definitions", 1)
    pri = doc.add_table(rows=1, cols=2)
    pri.style = "Table Grid"
    pri.rows[0].cells[0].text = "Priority"
    pri.rows[0].cells[1].text = "Definition"
    for p, d in [
        ("Critical", "Blocks launch — auth, onboarding, pay flow, safety, or data loss."),
        ("High", "Major feature broken; poor workaround."),
        ("Medium", "Secondary feature or UX issue with workaround."),
        ("Low", "Cosmetic or rare edge case."),
    ]:
        add_table_row(pri, [p, d])

    add_heading(doc, "5. How to Execute (Sequential)", 1)
    for line in [
        "Part A: Run TC-M-001 through TC-M-"
        f"{MOBILE_CASE_COUNT:03d} in order on iOS; repeat Critical + High on Android.",
        "Part B: Run TC-W-001 through TC-W-"
        f"{WEB_CASE_COUNT:03d} in order on desktop Chrome; spot-check Critical on Safari and mobile browser.",
        "Part C: After both parts, verify parity items — same Supabase data, consistent amounts and statuses.",
        "For each case record: Tester, Date, Environment (OS/build/browser), Result (Pass/Fail/Blocked), "
        "Actual result, Screenshot, Defect ID.",
        "Release gate: zero open Critical defects; all Critical and High cases Pass on both platforms.",
    ]:
        doc.add_paragraph(line, style="List Number")

    add_heading(doc, "6. Part A — LinkUp Mobile (Exclusive)", 1)
    add_para(
        doc,
        f"Execute sequentially: TC-M-001 to TC-M-{MOBILE_CASE_COUNT:03d}. "
        "Platform tag in preconditions: Mobile (iOS/Android app).",
    )
    mobile_written = render_suite(doc, MOBILE_SECTIONS, "TC-M", "Mobile (iOS/Android app)")

    add_heading(doc, "7. Part B — LinkUp Web (Exclusive)", 1)
    add_para(
        doc,
        f"Execute sequentially: TC-W-001 to TC-W-{WEB_CASE_COUNT:03d}. "
        "Platform tag in preconditions: Web (linkup-web browser).",
    )
    web_written = render_suite(doc, WEB_SECTIONS, "TC-W", "Web (linkup-web browser)")

    assert mobile_written == MOBILE_CASE_COUNT
    assert web_written == WEB_CASE_COUNT

    add_heading(doc, "8. Part C — Cross-Platform Parity Checklist", 1)
    add_para(doc, "Verify after Parts A and B; mark Pass/Fail for each row.")
    parity = doc.add_table(rows=1, cols=4)
    parity.style = "Table Grid"
    parity.rows[0].cells[0].text = "Flow"
    parity.rows[0].cells[1].text = "Mobile ref"
    parity.rows[0].cells[2].text = "Web ref"
    parity.rows[0].cells[3].text = "Result"
    parity_rows = [
        ("Signup → onboarding complete", "TC-M-002 area", "TC-W-002 area", "Pass / Fail"),
        ("Discover mood 20 km reach", "TC-M-003 area", "TC-W-003 area", "Pass / Fail"),
        ("Create paid plan → escrow funded", "TC-M-010 area", "TC-W-009 area", "Pass / Fail"),
        ("Dual meetup confirmation → disbursement", "TC-M-013 area", "TC-W-011 area", "Pass / Fail"),
        ("Group plan Annexure B lifecycle", "TC-M-014 area", "TC-W-012 area", "Pass / Fail"),
        ("Dispute + video evidence", "TC-M-016 area", "TC-W-013 area", "Pass / Fail"),
        ("Wallet balance after webhook", "TC-M-012 area", "TC-W-010 area", "Pass / Fail"),
        ("Notification deep link to plan/chat", "TC-M-020 area", "TC-W-016 area", "Pass / Fail"),
        ("KYC approve → verified badge", "TC-M-022 area", "TC-W-018 area", "Pass / Fail"),
        ("Admin resolve report/dispute", "TC-M-028 area", "TC-W-024 area", "Pass / Fail"),
    ]
    for row in parity_rows:
        add_table_row(parity, list(row))

    add_heading(doc, "9. Test Execution Summary", 1)
    summary = doc.add_table(rows=2, cols=7)
    summary.style = "Table Grid"
    headers = [
        "Mobile Total",
        "Mobile Passed",
        "Web Total",
        "Web Passed",
        "Combined Failed",
        "Blocked",
        "Pass Rate %",
    ]
    for i, h in enumerate(headers):
        summary.rows[0].cells[i].text = h
    summary.rows[1].cells[0].text = str(MOBILE_CASE_COUNT)
    summary.rows[1].cells[2].text = str(WEB_CASE_COUNT)

    add_heading(doc, "10. Launch Sign-off", 1)
    add_para(doc, "Sign only when Critical and High cases pass and no release-blocking defects remain open.")
    sign = doc.add_table(rows=5, cols=5)
    sign.style = "Table Grid"
    sign.rows[0].cells[0].text = "Role"
    sign.rows[0].cells[1].text = "Mobile UAT"
    sign.rows[0].cells[2].text = "Web UAT"
    sign.rows[0].cells[3].text = "Signature"
    sign.rows[0].cells[4].text = "Date"
    for role in ["QA Lead", "Product Owner", "Engineering Lead", "Release Manager"]:
        add_table_row(sign, [role, "", "", "", ""])

    return doc


def save_doc_from_docx(docx_path: Path, doc_path: Path) -> None:
    try:
        import shutil

        import win32com.client  # type: ignore

        temp_doc = doc_path.with_suffix(".generated.doc")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(temp_doc.resolve()), FileFormat=0)
        doc.Close(False)
        word.Quit()
        try:
            shutil.move(str(temp_doc), str(doc_path))
        except OSError:
            fallback = doc_path.with_name(doc_path.stem + "-generated.doc")
            shutil.move(str(temp_doc), str(fallback))
            print(f"Warning: {doc_path.name} is open or locked; wrote {fallback.name} instead.")
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"Could not write legacy .doc file: {exc}") from exc


def main() -> None:
    doc = build_document()
    import shutil

    temp_docx = DOCX_PATH.with_suffix(".generated.docx")
    doc.save(temp_docx)
    try:
        shutil.move(str(temp_docx), str(DOCX_PATH))
    except OSError:
        fallback_docx = DOCX_PATH.with_name(DOCX_PATH.stem + "-v2.0.docx")
        shutil.move(str(temp_docx), str(fallback_docx))
        print(f"Warning: {DOCX_PATH.name} is open; wrote {fallback_docx.name} instead.")
        save_doc_from_docx(fallback_docx, DOC_PATH.with_name(DOC_PATH.stem + "-v2.0.doc"))
        print(
            f"Wrote {fallback_docx.name} (Mobile: {MOBILE_CASE_COUNT}, "
            f"Web: {WEB_CASE_COUNT}, Total: {TOTAL_CASE_COUNT})"
        )
        return

    save_doc_from_docx(DOCX_PATH, DOC_PATH)
    print(
        f"Wrote {DOCX_PATH.name} and {DOC_PATH.name} "
        f"(Mobile: {MOBILE_CASE_COUNT}, Web: {WEB_CASE_COUNT}, Total: {TOTAL_CASE_COUNT})"
    )


if __name__ == "__main__":
    main()
