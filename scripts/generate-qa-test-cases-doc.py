"""Generate LinkUp QA user test cases Word document."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "LinkUp-QA-User-Test-Cases.docx"

PRIORITY = ("Critical", "High", "Medium", "Low")

# TC-ID, Module, Priority, Scenario, Preconditions, Steps, Expected Result
TEST_CASES = [
    # --- Authentication ---
    ("TC-AUTH-001", "Authentication", "Critical", "Cold start splash displays brand",
     "App not running; user logged out or first install.",
     "1. Launch LinkUp on device.\n2. Observe splash overlay.",
     "Pastel splash shows linkup logo lockup and “Meet With Confidence”; fades into login or app within ~5s after auth bootstrap."),
    ("TC-AUTH-002", "Authentication", "Critical", "Email signup with confirmation",
     "Valid unused email; Supabase confirm-email ON; SMTP configured.",
     "1. Open Sign up.\n2. Enter name, email, password.\n3. Accept privacy consent.\n4. Tap Sign up.",
     "Success message to check email; verification email received with tappable confirm link; no session until confirmed."),
    ("TC-AUTH-003", "Authentication", "Critical", "Email confirmation deep link",
     "Signup completed; confirmation email received on same device with app installed.",
     "1. Tap Confirm link in email.\n2. Observe app opens via linkup://auth/callback.",
     "App opens; session established; user routed to onboarding (new) or main tabs (returning)."),
    ("TC-AUTH-004", "Authentication", "High", "Resend verification email cooldown",
     "Signup awaiting confirmation.",
     "1. Tap Resend verification.\n2. Immediately tap Resend again.",
     "First resend succeeds or shows check email; second attempt blocked with cooldown message (~60s)."),
    ("TC-AUTH-005", "Authentication", "Critical", "Email/password login",
     "Confirmed account exists.",
     "1. Open Log in.\n2. Enter valid credentials.\n3. Tap Log in.",
     "User enters app; routed per onboarding_status (onboarding or Discover tab)."),
    ("TC-AUTH-006", "Authentication", "High", "Google OAuth sign-in",
     "Google provider enabled in Supabase; device has Google account.",
     "1. Tap Continue with Google.\n2. Complete Google consent.",
     "Returns to app with session; no separate email confirmation step."),
    ("TC-AUTH-007", "Authentication", "High", "Forgot password flow",
     "Confirmed account exists; reset email configured.",
     "1. Tap Forgot password.\n2. Enter email.\n3. Submit.\n4. Open reset link on same device.\n5. Set new password.",
     "Reset email received; deep link opens reset screen; new password saves; user can log in."),
    ("TC-AUTH-008", "Authentication", "Medium", "Show/hide password toggle",
     "On login or signup screen.",
     "1. Enter password.\n2. Tap eye icon to show.\n3. Tap again to hide.",
     "Password visibility toggles without clearing field."),
    ("TC-AUTH-009", "Authentication", "High", "Invalid login credentials",
     "Account exists.",
     "1. Enter wrong password.\n2. Tap Log in.",
     "Clear error message; user remains on login; no crash."),
    ("TC-AUTH-010", "Authentication", "Medium", "Signup without privacy consent",
     "On signup screen.",
     "1. Fill fields.\n2. Do not check privacy consent.\n3. Tap Sign up.",
     "Consent error shown; signup blocked."),
    # --- Onboarding ---
    ("TC-ONB-001", "Onboarding", "Critical", "Complete 5-step onboarding wizard",
     "New user; email confirmed; onboarding_status pending.",
     "1. Complete P1 Basics (name, DOB 18+, photos, intro video).\n2. Complete P2 Story.\n3. Complete P3 Location & prefs.\n4. Complete P4 Safety acknowledgment.\n5. Publish on P5 Preview.",
     "User lands on Discover tab; onboarding not shown again on relaunch."),
    ("TC-ONB-002", "Onboarding", "High", "Resume onboarding mid-wizard",
     "User partially completed onboarding; app killed mid-flow.",
     "1. Relaunch app.\n2. Continue from saved step.",
     "Wizard resumes at last saved step; prior entries preserved."),
    ("TC-ONB-003", "Onboarding", "High", "18+ age gate",
     "On P1 Basics.",
     "1. Enter birth date under 18.\n2. Attempt continue.",
     "Validation blocks progress with age requirement message."),
    ("TC-ONB-004", "Onboarding", "High", "Primary photo selection",
     "Multiple photos uploaded on P1.",
     "1. Set one photo as primary.\n2. Complete onboarding.\n3. View profile preview.",
     "Primary photo appears on profile card and discovery surfaces."),
    # --- Discover ---
    ("TC-DIS-001", "Discover", "Critical", "Discover tab loads swipe deck",
     "Onboarding complete; location optional.",
     "1. Open app (default Discover tab).\n2. Wait for feed load.",
     "Swipe cards or empty state with Create FAB; no infinite spinner."),
    ("TC-DIS-002", "Discover", "High", "Swipe right (interest) on plan",
     "Feed has at least one plan; user unverified OK for browse.",
     "1. Swipe right or tap like on a plan card.",
     "Interest path opens or saves interest per product rules; no crash."),
    ("TC-DIS-003", "Discover", "High", "Swipe left (hide) plan",
     "Premium user with undo available.",
     "1. Swipe left to hide plan.\n2. Tap Undo in header.",
     "Plan hidden then restored to deck."),
    ("TC-DIS-004", "Discover", "High", "Switch list display mode",
     "On Discover.",
     "1. Open Filters.\n2. Select List display.\n3. Apply.",
     "Feed shows scrollable list with search bar; preference persists on relaunch."),
    ("TC-DIS-005", "Discover", "Medium", "Apply feed filters",
     "Plans exist matching and not matching filters.",
     "1. Open Filters.\n2. Set distance, price, mood, verified hosts.\n3. Apply.",
     "Deck refreshes to match filters; filters persist in preferences."),
    ("TC-DIS-006", "Discover", "High", "Open plan detail from card",
     "Feed has plans.",
     "1. Tap plan card (or info action).",
     "Plan detail screen opens with host info, schedule, actions."),
    ("TC-DIS-007", "Discover", "High", "Open host member profile",
     "Plan card visible.",
     "1. Tap host avatar.",
     "Member profile (/user/[id]) opens with gallery and message option."),
    ("TC-DIS-008", "Discover", "High", "Unverified user sees KYC banner",
     "Onboarding done; KYC not approved.",
     "1. Open Discover.",
     "PlansKycBanner or soft KYC prompt visible; browsing still allowed."),
    ("TC-DIS-009", "Discover", "Critical", "Create plan hard gate (unverified)",
     "User unverified.",
     "1. Tap + Create FAB or create entry.",
     "Hard gate: verification required; option to start KYC or dismiss."),
    # --- Meetr ---
    ("TC-MET-001", "Meetr", "High", "Meetr grid loads meet types",
     "Onboarding complete.",
     "1. Navigate to Meetr tab.\n2. Observe category tiles.",
     "Portrait tiles show cover/gradient, meet type name (2 lines visible), no clipped text."),
    ("TC-MET-002", "Meetr", "High", "Browse plans by meet type",
     "Catalog meet type exists with plans.",
     "1. Tap a meet type tile.",
     "Discover or filtered view shows plans for that meet type."),
    ("TC-MET-003", "Meetr", "Medium", "Custom meet type pending state",
     "User submitted custom meet type awaiting approval.",
     "1. Open Meetr.\n2. Tap pending tile.",
     "Pending badge shown; modal explains approval status."),
    ("TC-MET-004", "Meetr", "Medium", "Press/hold reveals meet type description",
     "Meet type has description.",
     "1. Press and hold tile.\n2. Release.",
     "Description overlay expands; title hides; returns on release."),
    # --- Plans ---
    ("TC-PLN-001", "Plans", "Critical", "Verified user creates paid plan",
     "User KYC verified.",
     "1. Start Create plan.\n2. Complete meet/schedule, commitment, details.\n3. Publish.",
     "Success screen; plan appears on Discover for eligible viewers."),
    ("TC-PLN-002", "Plans", "High", "Create mood plan",
     "User verified.",
     "1. Create plan with mood/vibe fields.\n2. Publish.",
     "Plan appears in mood timeline on Discover."),
    ("TC-PLN-003", "Plans", "High", "Guest sends negotiation offer",
     "Verified guest; open plan accepting offers.",
     "1. Open plan detail.\n2. Open negotiate.\n3. Submit offer amount/message.",
     "Offer sent; appears in Offers tab (sent); host notified."),
    ("TC-PLN-004", "Plans", "Critical", "Unverified guest blocked from offer",
     "Guest not verified.",
     "1. Open plan.\n2. Attempt send offer.",
     "Verification gate shown; offer not sent."),
    ("TC-PLN-005", "Plans", "High", "Host accepts offer → agreement",
     "Host has pending acceptable offer.",
     "1. Open Offers or plan.\n2. Accept offer.\n3. Review agreement screen.",
     "Agreement summary shown; path to confirm or escrow for paid plans."),
    ("TC-PLN-006", "Plans", "Medium", "Edit published plan",
     "User is plan host.",
     "1. Open own plan.\n2. Open edit sheet.\n3. Change title/schedule.\n4. Save.",
     "Changes persist; feed reflects updates."),
    # --- Offers ---
    ("TC-OFR-001", "Offers", "High", "Offers tab sent/received lists",
     "User has sent and received offers.",
     "1. Open Offers tab.\n2. Switch sent/received segments.",
     "Correct offers listed with status (pending, countered, accepted)."),
    ("TC-OFR-002", "Offers", "High", "Counter-offer flow",
     "Pending offer exists.",
     "1. Open negotiation thread.\n2. Submit counter offer.",
     "Counter recorded; other party sees updated status."),
    # --- Escrow ---
    ("TC-ESC-001", "Escrow", "Critical", "Fund escrow via Paystack (paid plan)",
     "Agreement reached on paid plan; Paystack configured.",
     "1. Proceed to payment from agreement.\n2. Complete Paystack checkout.",
     "Escrow marked funded; plan active; both parties notified."),
    ("TC-ESC-002", "Escrow", "High", "View escrow status",
     "Active escrow exists.",
     "1. Open escrow detail from plan or wallet path.",
     "Status, amounts, counterparty, and next actions visible."),
    ("TC-ESC-003", "Escrow", "High", "Open dispute on escrow",
     "Eligible active escrow.",
     "1. Open dispute flow.\n2. Submit reason and evidence.",
     "Dispute created; funds flagged; admins notified."),
    # --- Messaging ---
    ("TC-MSG-001", "Messages", "Critical", "Messages inbox loads",
     "User has at least one conversation.",
     "1. Open Messages tab.",
     "Conversation list with preview, time, unread state."),
    ("TC-MSG-002", "Messages", "High", "Send text message in DM",
     "DM thread open.",
     "1. Type message.\n2. Send.",
     "Message appears in thread; delivered to recipient inbox (second device/account)."),
    ("TC-MSG-003", "Messages", "High", "Open chat from plan negotiation",
     "Active negotiation thread.",
     "1. Tap Open chat from negotiate screen.",
     "Chat opens with plan/offer context."),
    ("TC-MSG-004", "Messages", "Medium", "Smart suggestions bar",
     "Chat with plan/meetup context.",
     "1. Open chat thread.\n2. Observe suggestion chips.\n3. Tap a suggestion.",
     "Contextual suggestions shown; tap inserts/sends suggested text."),
    ("TC-MSG-005", "Messages", "High", "Group chat",
     "User member of group conversation.",
     "1. Open group thread.\n2. Send message.\n3. Open group info.",
     "Group messages deliver; member list accessible."),
    ("TC-MSG-006", "Messages", "Medium", "Push notification tap opens chat",
     "Push enabled; message received while app backgrounded.",
     "1. Receive push for new message.\n2. Tap notification.",
     "App opens correct chat thread."),
    # --- KYC ---
    ("TC-KYC-001", "Verification", "Critical", "Start KYC from hard gate",
     "Unverified user triggered gate (e.g. create plan).",
     "1. Tap Start verification.\n2. Begin K1 intro.",
     "KYC wizard opens; user can progress or exit with Not now."),
    ("TC-KYC-002", "Verification", "High", "Complete KYC document + video steps",
     "User in KYC wizard.",
     "1. Select ID type.\n2. Capture/upload ID.\n3. Record verification video.\n4. Submit.",
     "Status shows in progress; user notified on approval/rejection."),
    ("TC-KYC-003", "Verification", "High", "Verified user creates plan after approval",
     "KYC approved.",
     "1. Attempt create plan again.",
     "No hard gate; create wizard proceeds."),
    ("TC-KYC-004", "Verification", "Medium", "Soft KYC prompt dismiss",
     "Soft prompt shown on Discover.",
     "1. Tap Skip for now.",
     "Prompt dismisses; Discover usable; gate still applies on restricted actions."),
    # --- Notifications ---
    ("TC-NOT-001", "Notifications", "High", "Notification inbox loads",
     "User has system/message notifications.",
     "1. Open Notifications from bell/header entry.",
     "List grouped/filtered; unread indicators correct."),
    ("TC-NOT-002", "Notifications", "High", "Tap notification deep link",
     "Notification with plan/chat/admin target.",
     "1. Tap a notification row.",
     "Navigates to correct screen (plan, chat, admin tab, etc.)."),
    ("TC-NOT-003", "Notifications", "Medium", "Admin-only notification hidden from members",
     "Member account (non-admin); admin meet type notification exists.",
     "1. Open notification inbox as member.",
     "Admin-only types (e.g. meet type pending approval) not visible."),
    ("TC-NOT-004", "Notifications", "High", "Admin sees meet type pending notification",
     "Admin account; member submitted custom meet type.",
     "1. Open inbox as admin.\n2. Tap meet type notification.",
     "Notification visible; opens Admin → Meet types tab."),
    # --- Profile & Settings ---
    ("TC-SET-001", "Profile & Settings", "High", "Edit profile",
     "Logged in.",
     "1. Profile tab → Edit profile.\n2. Change bio/photos.\n3. Save.",
     "Changes visible on profile and discovery card."),
    ("TC-SET-002", "Profile & Settings", "High", "Notification preferences toggle",
     "On Settings → Notifications.",
     "1. Toggle Push off/on.\n2. Toggle Email off/on.\n3. Relaunch app.",
     "Preferences persist in profile preferences JSON."),
    ("TC-SET-003", "Profile & Settings", "Medium", "Travel mode (Premium)",
     "Premium user.",
     "1. Settings → Travel.\n2. Set travel location.\n3. Open Discover header.",
     "Travel label shown; feed ranking uses travel location."),
    ("TC-SET-004", "Profile & Settings", "High", "Delete account flow",
     "Logged in; no blocking legal hold.",
     "1. Settings → Delete account.\n2. Confirm deletion.",
     "Account deletion initiated/completed per policy; user signed out."),
    ("TC-SET-005", "Profile & Settings", "Medium", "Privacy policy re-consent",
     "Policy update required flag set.",
     "1. Launch app.\n2. Complete re-consent screen.",
     "Access restored after accept; decline blocks or limits per policy."),
    # --- Premium & Wallet ---
    ("TC-PRM-001", "Premium", "High", "Premium subscription checkout",
     "Paystack premium configured.",
     "1. Profile → Premium.\n2. Select tier.\n3. Complete checkout.",
     "Subscription active; badge/perks visible."),
    ("TC-PRM-002", "Premium", "Medium", "Saved plans tab (Premium)",
     "Premium user; saved plans exist.",
     "1. Open Saved tab.",
     "Bookmarked plans listed."),
    ("TC-WAL-001", "Wallet", "High", "Wallet balance and ledger",
     "User with transaction history.",
     "1. Open Wallet tab.",
     "Balance, goodwill credits, ledger entries visible and accurate."),
    # --- Support ---
    ("TC-SUP-001", "Support", "High", "Create support ticket",
     "Logged in.",
     "1. Open Support.\n2. Create ticket with subject/body.\n3. Submit.",
     "Ticket created; appears in user's ticket list."),
    ("TC-SUP-002", "Support", "Medium", "Reply on existing ticket",
     "Open ticket exists.",
     "1. Open ticket detail.\n2. Send reply.",
     "Reply appended; status updated."),
    ("TC-DSP-001", "Disputes", "High", "View disputes list",
     "User has filed or received dispute.",
     "1. Open Disputes from profile/support path.",
     "Disputes listed with status and plan reference."),
    # --- Admin ---
    ("TC-ADM-001", "Admin", "High", "Admin dashboard access control",
     "Non-admin and admin accounts available.",
     "1. Open /admin as non-admin.\n2. Open as admin.",
     "Non-admin blocked; admin sees queues (KYC, reports, meet types, etc.)."),
    ("TC-ADM-002", "Admin", "High", "Approve custom meet type",
     "Pending user meet type in queue.",
     "1. Admin → Meet types.\n2. Approve type.",
     "Type active for users; submitter notified."),
    ("TC-ADM-003", "Admin", "High", "KYC approval grants Silver trial",
     "Pending KYC submission.",
     "1. Admin KYC queue → Approve.",
     "User verified; trial/subscription event per product rules."),
    # --- Regression / Non-functional ---
    ("TC-NFR-001", "Non-functional", "High", "App background and resume",
     "User logged in mid-flow (chat or discover).",
     "1. Background app 2 min.\n2. Resume.",
     "Session intact; screen restores without crash."),
    ("TC-NFR-002", "Non-functional", "Medium", "Pull to refresh on Meetr/Discover",
     "On Meetr or Discover with network.",
     "1. Pull to refresh.",
     "Content reloads; spinner clears."),
    ("TC-NFR-003", "Non-functional", "High", "Offline graceful error",
     "Enable airplane mode.",
     "1. Open Discover or Messages.",
     "Friendly error/empty state; no white screen crash."),
]


def add_title_page(doc: Document) -> None:
    title = doc.add_heading("LinkUp — User Test Cases", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Quality Assurance · Manual User Acceptance Testing")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Document version: 1.0\nGenerated: {date.today().isoformat()}\nProduct: LinkUp Mobile (iOS / Android)\n").font.size = Pt(11)

    doc.add_page_break()


def add_intro(doc: Document) -> None:
    doc.add_heading("1. Purpose & Scope", level=1)
    doc.add_paragraph(
        "This document defines manual user test cases for LinkUp — a social meetup app with "
        "discovery, plans, negotiation, escrow, messaging, verification (KYC), premium features, "
        "and admin tooling. Cases are written for QA testers, UAT participants, and release sign-off."
    )

    doc.add_heading("2. Test Environment", level=1)
    items = [
        "Staging or production Supabase project with Auth SMTP configured (signup/reset emails deliver).",
        "Physical device recommended for push notifications, deep links (linkup://auth/callback), and KYC camera flows.",
        "Test accounts: unverified member, verified member, premium member, admin (admins table row).",
        "Paystack test keys for escrow and subscription flows.",
        "Optional second device/account for messaging and offer negotiation.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Priority Definitions", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    headers = ("Priority", "Definition")
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ("Critical", "Blocks release; core auth, onboarding, discover, plan pay flow, or data loss risk."),
        ("High", "Major feature broken; workaround difficult."),
        ("Medium", "Secondary feature or UX issue with workaround."),
        ("Low", "Cosmetic or edge case."),
    ]
    for r, (p, d) in enumerate(rows, start=1):
        table.rows[r].cells[0].text = p
        table.rows[r].cells[1].text = d

    doc.add_heading("4. How to Execute", level=1)
    doc.add_paragraph(
        "For each test case record: Tester name, Date, Environment (OS/build), Result (Pass/Fail/Blocked), "
        "Actual result, Screenshot/reference, Defect ID if failed."
    )
    doc.add_page_break()


def add_test_case_section(doc: Document) -> None:
    doc.add_heading("5. Test Cases", level=1)

    current_module = None
    for tc in TEST_CASES:
        tc_id, module, priority, scenario, pre, steps, expected = tc
        if module != current_module:
            current_module = module
            doc.add_heading(module, level=2)

        doc.add_heading(f"{tc_id}: {scenario}", level=3)

        meta = doc.add_table(rows=2, cols=4)
        meta.style = "Table Grid"
        labels = [("Test ID", tc_id), ("Module", module), ("Priority", priority), ("Result", "Pass / Fail / Blocked")]
        for c, (label, value) in enumerate(labels):
            meta.rows[0].cells[c].text = label
            meta.rows[1].cells[c].text = value

        doc.add_paragraph("Preconditions:", style="Heading 4")
        doc.add_paragraph(pre)

        doc.add_paragraph("Test Steps:", style="Heading 4")
        for line in steps.split("\n"):
            doc.add_paragraph(line, style="List Number")

        doc.add_paragraph("Expected Result:", style="Heading 4")
        p = doc.add_paragraph(expected)
        p.runs[0].bold = False

        doc.add_paragraph("")  # spacer


def add_signoff(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("6. Test Execution Summary", level=1)
    summary = doc.add_table(rows=2, cols=6)
    summary.style = "Table Grid"
    cols = ["Total Cases", "Passed", "Failed", "Blocked", "Not Run", "Pass Rate %"]
    for i, c in enumerate(cols):
        summary.rows[0].cells[i].text = c
    summary.rows[1].cells[0].text = str(len(TEST_CASES))

    doc.add_paragraph("")
    doc.add_heading("7. Sign-off", level=1)
    sign = doc.add_table(rows=4, cols=4)
    sign.style = "Table Grid"
    for i, h in enumerate(["Role", "Name", "Signature", "Date"]):
        sign.rows[0].cells[i].text = h
    for i, role in enumerate(["QA Lead", "Product Owner", "Engineering Lead"], start=1):
        sign.rows[i].cells[0].text = role


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    add_title_page(doc)
    add_intro(doc)
    add_test_case_section(doc)
    add_signoff(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({len(TEST_CASES)} test cases)")


if __name__ == "__main__":
    main()
